import os
import tempfile
from docx import Document
from flask import send_file

def get_disease_info(disease):
    descriptions = {
        "Early Blight": (
            "Early blight is a common fungal disease that affects tomato and potato crops. It is caused by the pathogen *Alternaria solani* and is characterized by dark, concentric spots on lower, older leaves. As the disease progresses, leaves turn yellow and drop prematurely, reducing the plant's photosynthetic capacity and yield. Conditions that favor early blight include high humidity, warm temperatures, and prolonged leaf wetness. It is important to manage early blight promptly, as the disease can quickly spread and cause significant damage if left untreated. Cultural practices like crop rotation, resistant varieties, and proper spacing can reduce the risk of infection. Regular field inspections and early detection are key. Fungicide applications should be considered as part of an integrated pest management (IPM) strategy. Proper sanitation, such as removing crop debris, also helps in controlling the spread of the fungus. Healthy, vigorous plants are generally more resistant, so ensure optimal fertilization and irrigation practices."
        ),
        "Late Blight": (
            "Late blight is a devastating disease that affects tomato and potato plants. It is caused by the water mold *Phytophthora infestans*, the same pathogen responsible for the Irish Potato Famine. Symptoms include water-soaked lesions on leaves, stems, and fruits that rapidly enlarge and turn brown or black. In humid conditions, a white moldy growth may appear on the undersides of leaves. The disease spreads rapidly and can destroy entire fields in a matter of days. Late blight thrives in cool, moist environments, and can overwinter in infected tubers. Effective control involves planting disease-free seeds, avoiding overhead irrigation, and removing infected plant materials. Fungicides can provide control but must be applied preventatively. Integrated disease management practices are essential to limit its spread and impact."
        )
    }
    return descriptions.get(disease, "No description available for this disease.")

def generate_precaution_doc(disease, precaution):
    doc = Document()

    # Add the title
    doc.add_heading(f'Precaution Information for {disease}', 0)

    # Add description/introduction paragraph
    description = get_disease_info(disease)
    doc.add_paragraph(description)

    # Add Diagnosis Summary Table
    doc.add_heading('Diagnosis Summary:', level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'LightShading-Accent1'
    table.cell(0, 0).text = 'Category'
    table.cell(0, 1).text = 'Details'
    table.cell(1, 0).text = 'Disease'
    table.cell(1, 1).text = disease
    table.cell(2, 0).text = 'Confidence'
    table.cell(2, 1).text = '26.33%'  # Placeholder
    table.cell(3, 0).text = 'Symptoms'
    table.cell(3, 1).text = 'Refer to the introduction for key symptoms.'

    # Add diagnosis process section
    doc.add_heading('Diagnosis Stages of Tomato Blight:', level=1)
    diagnosis_steps = [
        "**1. Visual Inspection of Leaves:** Start by checking the lower (older) leaves of the plant. Early blight typically begins here. Look for irregular, brown spots with concentric rings that resemble a target or bullseye.",
        "**2. Observe Leaf Yellowing:** Surrounding the dark spots, you may notice yellowing of the leaf tissue. As the infection spreads, the entire leaf may turn yellow and eventually die.",
        "**3. Check for Leaf Drop:** Infected leaves may fall off prematurely. Excessive leaf drop is a major symptom of early blight and can reduce photosynthesis and yield.",
        "**4. Inspect Stems and Petioles:** Look for elongated, sunken, dark lesions on the stem and leaf stalks (petioles), especially near the soil line. These lesions may also have concentric ring patterns.",
        "**5. Examine Fruit (if present):** If the plant has started fruiting, check the fruit near the stem for dark, sunken spots. These are typically less common but indicate disease progression.",
        "**6. Assess Environmental Conditions:** Early blight thrives in warm (24–29°C), humid environments. Frequent dew, rain, or overhead irrigation can accelerate the disease. If these conditions are present, chances of early blight are higher.",
    ]
    for step in diagnosis_steps:
        doc.add_paragraph(step, style='List Bullet')

    # Precautionary Measures
    doc.add_heading('Precautionary Measures:', level=1)
    doc.add_paragraph(precaution)

    # Recommendations
    doc.add_heading('Recommendations:', level=1)
    doc.add_paragraph('1. Follow proper hygiene and crop rotation practices.')
    doc.add_paragraph('2. Regularly monitor your crops for early signs of infection.')

    # Solutions
    doc.add_heading('Solutions:', level=1)
    doc.add_paragraph('1. Apply the recommended fungicides as per the instructions.')
    doc.add_paragraph('2. Remove affected leaves and improve field sanitation.')

    # Save to temporary file
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, "precaution_info.docx")
    doc.save(file_path)
    return file_path

def download_precaution_doc(disease, precaution):
    file_path = generate_precaution_doc(disease, precaution)
    return send_file(file_path, as_attachment=True)
